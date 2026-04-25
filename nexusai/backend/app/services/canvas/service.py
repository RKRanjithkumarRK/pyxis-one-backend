"""Canvas service — AI edits and document export."""
from __future__ import annotations
import io
import logging
import re
from typing import Literal

import html2text
import litellm
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

logger = logging.getLogger("nexusai.canvas")

ExportFormat = Literal["md", "html", "docx"]


def _strip_html(html: str) -> str:
    """Convert HTML to plain text."""
    converter = html2text.HTML2Text()
    converter.ignore_links = False
    converter.ignore_images = True
    converter.body_width = 0
    return converter.handle(html).strip()


async def ai_edit(selected_text: str, instruction: str, *, context: str = "") -> str:
    """
    Rewrite `selected_text` according to `instruction`.
    Returns the rewritten text.
    """
    system = (
        "You are an expert writing assistant. Rewrite the provided text according to the instruction. "
        "Return ONLY the rewritten text — no explanations, no preamble, no quotes. "
        "Preserve the original formatting style (markdown, lists, headings) unless instructed otherwise."
    )
    context_block = f"\n\nDocument context (surrounding text, do not include in output):\n{context[:1000]}" if context else ""
    user = (
        f"Instruction: {instruction}\n\n"
        f"Text to rewrite:\n{selected_text}"
        f"{context_block}"
    )
    try:
        resp = await litellm.acompletion(
            model="claude-sonnet-4-20250514",
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            max_tokens=4096,
            temperature=0.5,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as exc:
        logger.warning("AI edit failed: %s", exc)
        raise


def export_docx(title: str, html_content: str) -> bytes:
    """Convert HTML content to a DOCX file and return bytes."""
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.style.font.size = Pt(20)

    soup = BeautifulSoup(html_content, "html.parser")

    def _process_node(node, doc: Document) -> None:
        if node.name in ("h1", "h2", "h3", "h4"):
            level = int(node.name[1])
            doc.add_heading(node.get_text(strip=True), level=level)
        elif node.name == "p":
            text = node.get_text(separator=" ", strip=True)
            if text:
                p = doc.add_paragraph(text)
                for child in node.children:
                    if hasattr(child, "name"):
                        if child.name == "strong":
                            for run in p.runs:
                                run.bold = True
                        elif child.name == "em":
                            for run in p.runs:
                                run.italic = True
        elif node.name in ("ul", "ol"):
            for li in node.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(separator=" ", strip=True), style="List Bullet")
        elif node.name == "blockquote":
            text = node.get_text(separator=" ", strip=True)
            if text:
                p = doc.add_paragraph(text)
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else p.style
        elif node.name == "pre":
            code = node.get_text()
            p = doc.add_paragraph(code)
            p.style = "No Spacing"
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        elif node.name == "hr":
            doc.add_paragraph("─" * 40)

    for child in soup.body.children if soup.body else soup.children:
        if hasattr(child, "name") and child.name:
            _process_node(child, doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_markdown(html_content: str) -> str:
    """Convert HTML to Markdown."""
    converter = html2text.HTML2Text()
    converter.ignore_links = False
    converter.ignore_images = False
    converter.body_width = 0
    converter.protect_links = True
    return converter.handle(html_content).strip()


def export_html(title: str, html_content: str) -> str:
    """Return a complete HTML document."""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; line-height: 1.6; color: #1a1a1a; }}
    h1 {{ font-size: 2rem; margin-bottom: 0.5rem; }}
    h2 {{ font-size: 1.5rem; margin-top: 2rem; }}
    h3 {{ font-size: 1.25rem; margin-top: 1.5rem; }}
    pre {{ background: #f4f4f5; padding: 1rem; border-radius: 0.5rem; overflow-x: auto; }}
    code {{ font-family: 'Fira Code', 'Cascadia Code', monospace; font-size: 0.875rem; }}
    blockquote {{ border-left: 3px solid #e5e7eb; margin: 1rem 0; padding-left: 1rem; color: #6b7280; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #e5e7eb; padding: 0.5rem 0.75rem; text-align: left; }}
    th {{ background: #f9fafb; font-weight: 600; }}
    a {{ color: #2563eb; }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  {html_content}
</body>
</html>"""
